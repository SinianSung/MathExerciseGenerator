import sys, getopt, random, math, datetime, time
from docx import *

def getParameters(argv):
    try:
        opts, args = getopt.getopt(argv,"hn:t:r:T:",["help","number=","types=","Title="])
    except getopt.GetoptError:
        print("use exerc_equatioins -n 4 -t beltqspr -T 'Kurztest 2'")
        sys.exit(2)
    for opt, arg in opts:
        if opt =='-h':
            print("use exerc_equatioins -n 4 -t beltqspr -T 'Kurztest 2'")
            sys.exit()
        elif opt in ("-n","--number"):
            number_of_equations = int(arg)
        elif opt in ("-t","--types"):
            list_of_types = list(arg)
        elif opt in ("-r","--random"):
            is_random = arg.lower() == 'true'
        elif opt in ("-T","--Title"):
            title = arg
        else:
            sys.exit()
        
    return [number_of_equations, list_of_types, is_random, title]


def check_types(t, lot):
    for item in lot:
        if item not in t:
            lot.remove(item)

    return lot

def createFile( funs,name ="Aufgaben", intro = "Solve"):
    """
    creates word document , exercises is a list of (equation, solution) pairs
    .isocalendar()[1]
    """
    #weeknumber = datetime.datetime.fromtimestamp(time.time()).isocalendar()[1]
    timestmp = datetime.date.fromtimestamp(time.time())
    document = Document()
    document.add_heading('Übungsaufgaben', 0)
    document.add_paragraph(intro)

    document.add_heading('Aufgaben', level=1)
    for item in funs:
        document.add_paragraph("{},".format(item[0]), style='List Number')

    document.add_page_break()

    document.add_heading('Lösungen', level=1)
    for item in funs:
        document.add_paragraph("{} ".format(item[1]),style='List Number 2')

    document.save('{}_{}.docx'.format(name, timestmp))


