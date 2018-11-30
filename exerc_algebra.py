import datetime
import math
import random
import time

from docx import *
from termGenerator import *


def createFile( exercises,name ="Term-Aufgaben"):
    """
    creates word document , exercises is a list of (term, solution) pairs
    .isocalendar()[1]
    """
    #weeknumber = datetime.datetime.fromtimestamp(time.time()).isocalendar()[1]
    timestmp = datetime.date.fromtimestamp(time.time())
    document = Document()
    document.add_heading('Aufgabenchallenge', 0)
    document.add_paragraph('Lösen Sie so viele Aufgaben wie möglich in 5 Minuten.')

    document.add_heading('Aufgaben', level=1)
    for item in exercises:
        document.add_paragraph("{},".format(item[0]), style='List Number')

    document.add_page_break()

    document.add_heading('Lösungen', level=1)
    for item in exercises:
        document.add_paragraph("{} ".format(item[1]),style='List Number 2')

    document.save('{}_{}.docx'.format(name, timestmp))

def main():
    number_of_terms = 60
    exercises = []
    types ={'fo': factor_out, \
    'mult':mult ,\
    'pm': prod_mult, \
    's2f': sum_to_factor, \
    'b2f': binom_to_factor, \
    'bm': binom_mult, \
    'bms': binom_mult_s, \
    'tri3': trinom_s,\
    'tri': trinom,\
    'bm3': binom_mult_3,\
    'term': termprodukt
        }

    for i in range(number_of_terms):
        choice = random.choice(['tri','bms','bm3','b2f']) # ['fo','mult','pm','s2f','b2f','bm']:
        exerc = types[choice]()
        exercises.append(exerc)
        exerc =""

    createFile(exercises,"Übungsaufgaben")  

if __name__ == "__main__":
    main()
